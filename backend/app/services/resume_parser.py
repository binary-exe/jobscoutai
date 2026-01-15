"""
Resume file parser for PDF and DOCX files.

Extracts text from uploaded resume files.
"""

from typing import Dict, Optional
from io import BytesIO
import re

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


async def parse_resume_file(file_content: bytes, filename: str) -> Dict[str, str]:
    """
    Parse a resume file (PDF or DOCX) and extract text.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename (for determining type)
    
    Returns:
        {
            "text": str,  # Extracted text
            "error": Optional[str],  # Error message if parsing failed
        }
    """
    file_lower = filename.lower()
    
    # Determine file type
    if file_lower.endswith('.pdf'):
        return await _parse_pdf(file_content)
    elif file_lower.endswith('.docx') or file_lower.endswith('.doc'):
        return await _parse_docx(file_content)
    else:
        return {
            "text": "",
            "error": f"Unsupported file type: {filename}. Supported: PDF, DOCX"
        }


async def _parse_pdf(file_content: bytes) -> Dict[str, str]:
    """Parse PDF file and extract text."""
    if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        return {
            "text": "",
            "error": "PDF parsing libraries not installed. Install pdfplumber or PyPDF2."
        }
    
    try:
        # Try pdfplumber first (better text extraction)
        if PDFPLUMBER_AVAILABLE:
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                text = "\n\n".join(text_parts)
                return {"text": _clean_text(text), "error": None}
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            pdf_reader = PdfReader(BytesIO(file_content))
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            text = "\n\n".join(text_parts)
            return {"text": _clean_text(text), "error": None}
        
        return {"text": "", "error": "No PDF parser available"}
        
    except Exception as e:
        return {
            "text": "",
            "error": f"Failed to parse PDF: {str(e)}"
        }


async def _parse_docx(file_content: bytes) -> Dict[str, str]:
    """Parse DOCX file and extract text."""
    if not DOCX_AVAILABLE:
        return {
            "text": "",
            "error": "DOCX parsing library not installed. Install python-docx."
        }
    
    try:
        doc = Document(BytesIO(file_content))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n\n".join(text_parts)
        return {"text": _clean_text(text), "error": None}
        
    except Exception as e:
        return {
            "text": "",
            "error": f"Failed to parse DOCX: {str(e)}"
        }


def _clean_text(text: str) -> str:
    """Clean extracted text."""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove special characters that might cause issues
    text = text.replace('\x00', '')  # Remove null bytes
    
    return text.strip()
