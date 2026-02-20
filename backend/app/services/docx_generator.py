"""
DOCX generation service for Apply Workspace.

Generates downloadable DOCX files for tailored resume and cover note.
Creates professional, ATS-friendly documents with proper formatting.
"""

from typing import Dict, Optional, Any, List, Tuple
from io import BytesIO
import re
import zipfile
from urllib.parse import urlparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ==================== Resume Section Parsing ====================

_BULLET_PREFIX_RE = re.compile(
    r"^\s*(?:[•\-\*\u2022\u25cf\u25aa]|(?:\d+[\.\)])|(?:[a-zA-Z][\.\)]))\s+"
)

_DATE_RANGE_ONLY_RE = re.compile(
    r"^\s*(?:\w+\s+20\d{2}|\b20\d{2})\s*[–—-]\s*(?:\w+\s+20\d{2}|Present|\b20\d{2})\s*$",
    re.IGNORECASE,
)


def _strip_bullet_prefix(s: str) -> str:
    t = (s or "").strip()
    return _BULLET_PREFIX_RE.sub("", t).strip()


def _strip_leading_markers(s: str) -> str:
    # Resume text often contains decorative glyphs; strip them for ATS/consistency.
    return (s or "").lstrip("◆■●►⚫★▶━═—⧉•*- ").strip()


def _is_probable_name_line(line: str) -> bool:
    """
    Heuristic: pick the first "name-like" line in the header area.
    Avoid contact lines, URLs, and lines heavy with digits.
    """
    t = (line or "").strip()
    if not t:
        return False
    if _is_contact_line(t) or _is_section_header(t):
        return False
    if re.search(r"\d", t):
        return False
    words = [w for w in re.split(r"\s+", t) if w]
    if len(words) < 2 or len(words) > 5:
        return False
    if sum(1 for ch in t if ch.isalpha()) < 6:
        return False
    # Reject "role • role • role" headlines
    if "•" in t and len(words) > 3:
        return False
    return True

def _parse_resume_into_structure(resume_text: str) -> Dict[str, Any]:
    """
    Parse resume text into a structured format with proper sections.
    Returns a dict with: header, contact, summary, skills, experience, projects, education, certifications, etc.
    """
    result = {
        "name": "",
        "title": "",
        "location": "",
        "contact": [],  # List of {type, value, url}
        "summary": "",
        "skills": [],  # List of skill strings or {category: [skills]}
        "experience": [],  # List of {title, company, location, dates, bullets}
        "projects": [],  # List of {name, description, url, bullets}
        "education": [],  # List of {degree, school, dates}
        "certifications": [],  # List of {name, issuer, date}
        "achievements": [],  # List of strings
        "languages": [],
        "other_sections": {},  # Any unrecognized sections
    }
    
    lines = resume_text.strip().split("\n")
    if not lines:
        return result

    # Find where sections start (after header/contact area)
    section_start_idx = 0
    for i, line in enumerate(lines):
        if _is_section_header(line):
            section_start_idx = i
            break

    header_lines = [ln.strip() for ln in lines[: max(0, section_start_idx)] if ln.strip()]

    # Name: pick first probable name line; otherwise fall back to first non-contact line.
    name = ""
    for ln in header_lines[:6]:
        if _is_probable_name_line(ln):
            name = ln.strip()
            break
    if not name:
        for ln in header_lines[:6]:
            if not _is_contact_line(ln) and not _is_section_header(ln):
                name = ln.strip()
                break
    result["name"] = name

    # Title: next non-contact, non-section line after name (if any).
    title = ""
    if name:
        after_name = False
        for ln in header_lines[:10]:
            if ln.strip() == name:
                after_name = True
                continue
            if not after_name:
                continue
            if not _is_contact_line(ln) and not _is_section_header(ln) and not _looks_like_location(ln):
                title = ln.strip()
                break
    result["title"] = title

    # Contact/location: scan the header region (all lines before first section header)
    for ln in header_lines[:12]:
        if _looks_like_location(ln) and not result["location"]:
            result["location"] = ln.strip()
        if _is_contact_line(ln):
            result["contact"].extend(_extract_contact_items(ln))

    # De-dupe contacts by (type,value,url)
    deduped: List[Dict[str, str]] = []
    seen = set()
    for c in result["contact"]:
        key = (c.get("type"), c.get("value"), c.get("url"))
        if key in seen:
            continue
        seen.add(key)
        deduped.append(c)
    result["contact"] = deduped
    
    # Now parse sections
    current_section = None
    current_content: List[str] = []
    
    # Parse sections
    for i, line in enumerate(lines[section_start_idx:], section_start_idx):
        line_stripped = line.strip()
        
        if _is_section_header(line_stripped):
            # Save previous section
            if current_section:
                _save_section(result, current_section, current_content)
            current_section = _normalize_section_name(line_stripped)
            current_content = []
        elif current_section:
            current_content.append(line_stripped)
    
    # Save last section
    if current_section:
        _save_section(result, current_section, current_content)
    
    return result


def _is_contact_line(line: str) -> bool:
    """Check if line contains contact information."""
    contact_patterns = [
        r'@', r'linkedin\.com', r'github\.com', r'upwork\.com',
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # Phone
        r'http[s]?://', r'www\.',
    ]
    line_lower = line.lower()
    return any(re.search(p, line_lower) for p in contact_patterns)


def _looks_like_location(line: str) -> bool:
    """Check if line looks like a location."""
    location_keywords = ['remote', 'dubai', 'new york', 'london', 'berlin', 'singapore', 
                        'united states', 'united kingdom', 'uae', 'emirates', 'india']
    line_lower = line.lower()
    return any(kw in line_lower for kw in location_keywords) and len(line) < 100


def _extract_contact_items(line: str) -> List[Dict[str, str]]:
    """Extract contact items from a line."""
    items = []

    # Emails (can be multiple)
    for email in re.findall(r"[\w\.-]+@[\w\.-]+\.\w+", line):
        items.append({"type": "email", "value": email, "url": f"mailto:{email}"})

    # Phones (can be multiple)
    for phone in re.findall(r"\+?\d[\d\s\-\.]{8,}\d", line):
        items.append({"type": "phone", "value": phone.strip()})

    # LinkedIn (with or without scheme)
    for m in re.findall(
        r"(https?://(?:www\.)?linkedin\.com/[^\s|]+|(?:www\.)?linkedin\.com/[^\s|]+|linkedin\.com/[^\s|]+)",
        line,
        flags=re.I,
    ):
        url = m.strip()
        if not url.startswith("http"):
            url = "https://" + url.lstrip("/")
        items.append({"type": "linkedin", "value": "LinkedIn", "url": url})

    # GitHub (with or without scheme)
    for m in re.findall(
        r"(https?://(?:www\.)?github\.com/[A-Za-z0-9_-]+|(?:www\.)?github\.com/[A-Za-z0-9_-]+|github\.com/[A-Za-z0-9_-]+)",
        line,
        flags=re.I,
    ):
        url = m.strip()
        if not url.startswith("http"):
            url = "https://" + url.lstrip("/")
        items.append({"type": "github", "value": "GitHub", "url": url})

    def _is_valid_http_url(u: str) -> bool:
        try:
            p = urlparse(u)
            if p.scheme not in ("http", "https"):
                return False
            host = (p.netloc or "").lower()
            if not host or "." not in host:
                return False
            # Reject "www.something" without a real TLD (common extraction artifact)
            if host.startswith("www.") and host.count(".") == 1:
                return False
            # Require a plausible TLD
            if not re.search(r"\.[a-z]{2,}$", host):
                return False
            return True
        except Exception:
            return False

    # Other URLs
    for url in re.findall(r"https?://[^\s|]+", line, flags=re.I):
        if "linkedin.com" in url.lower() or "github.com" in url.lower():
            continue
        if _is_valid_http_url(url):
            items.append({"type": "website", "value": url, "url": url})

    # Bare domains (portfolio) e.g. abdullahjavaid.me
    # Keep strict to avoid broken matches like "https://www.upwork"
    for domain in re.findall(r"\b([A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+)\b", line):
        d = domain.lower().strip(".")
        if any(x in d for x in ("linkedin.com", "github.com", "upwork.com")):
            continue
        if d.startswith("www.") and d.count(".") == 1:
            continue
        if d in ("gmail.com", "outlook.com", "hotmail.com"):
            continue
        # Only accept plausible domains with a TLD of 2+ chars
        if not re.search(r"\.[a-z]{2,}$", d):
            continue
        items.append({"type": "website", "value": domain, "url": "https://" + domain})
    
    return items


def _is_section_header(line: str) -> bool:
    """Check if line is a section header."""
    line = line.strip()
    if not line:
        return False
    
    # Check for common section markers
    section_markers = ['◆', '■', '●', '►', '⚫', '★', '▶', '━', '═', '—']
    if any(line.startswith(m) for m in section_markers):
        return True
    
    # Check for common section names
    section_keywords = [
        'profile', 'summary', 'objective', 'about',
        'experience', 'employment', 'work history', 'career',
        'education', 'academic', 'qualification',
        'skills', 'technical', 'competencies', 'expertise',
        'projects', 'portfolio',
        'certifications', 'certificates', 'licenses',
        'achievements', 'accomplishments', 'awards', 'honors',
        'languages', 'interests', 'hobbies',
    ]
    line_lower = line.lower().strip('◆■●►⚫★▶━═—: ')
    return any(kw in line_lower for kw in section_keywords) and len(line) < 50


def _normalize_section_name(line: str) -> str:
    """Normalize section header to standard name."""
    line_lower = line.lower().strip('◆■●►⚫★▶━═—: ')
    
    if any(kw in line_lower for kw in ['profile', 'summary', 'objective', 'about']):
        return 'summary'
    if any(kw in line_lower for kw in ['experience', 'employment', 'work', 'career']):
        return 'experience'
    if any(kw in line_lower for kw in ['education', 'academic', 'qualification']):
        return 'education'
    if any(kw in line_lower for kw in ['skill', 'technical', 'competenc', 'expertise', 'stack']):
        return 'skills'
    if any(kw in line_lower for kw in ['project', 'portfolio']):
        return 'projects'
    if any(kw in line_lower for kw in ['certification', 'certificate', 'license']):
        return 'certifications'
    if any(kw in line_lower for kw in ['achievement', 'accomplishment', 'award', 'honor', 'notable']):
        return 'achievements'
    if 'language' in line_lower:
        return 'languages'
    return 'other'


def _save_section(result: Dict, section_name: str, content: List[str]) -> None:
    """Save parsed section content to result dict."""
    content_text = '\n'.join(content).strip()
    if not content_text:
        return
    
    if section_name == 'summary':
        result['summary'] = content_text
    elif section_name == 'skills':
        result['skills'] = _parse_skills(content)
    elif section_name == 'experience':
        result['experience'] = _parse_experience(content)
    elif section_name == 'projects':
        result['projects'] = _parse_projects(content)
    elif section_name == 'education':
        edu, cert_like = _parse_education(content)
        result['education'] = edu
        if cert_like:
            # Best-effort: treat "Issued ..." lines as certifications, rather than education.
            result.setdefault("certifications", [])
            result["certifications"].extend(cert_like)
    elif section_name == 'certifications':
        result['certifications'] = _parse_certifications(content)
    elif section_name == 'achievements':
        result['achievements'] = _parse_bullets(content)
    elif section_name == 'languages':
        result['languages'] = _parse_bullets(content)
    else:
        result['other_sections'][section_name] = content_text


def _parse_skills(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse skills section into categorized skills."""
    skills = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('—') or line.startswith('━'):
            continue
        if line.lower().startswith("relevant keywords:"):
            continue
        # Check if it's a categorized skill line (Category: skill1, skill2)
        if ':' in line:
            parts = line.split(':', 1)
            category = parts[0].strip('•-* ')
            skill_list = [s.strip() for s in parts[1].split(',')]
            skills.append({"category": category, "items": skill_list})
        else:
            # Single skill or bullet point
            skill = line.strip('•-* ')
            if skill:
                skills.append({"category": None, "items": [skill]})
    return skills


def _parse_experience(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse experience section into structured entries."""
    experiences = []
    current_exp = None
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith('—') or line.startswith('━'):
            continue
        
        # Bullet lines should never start a new experience entry
        is_bullet_line = bool(_BULLET_PREFIX_RE.match(line)) or line.startswith("•")

        # If a date-range appears on its own line, attach it to the current experience.
        if current_exp and (not is_bullet_line) and _DATE_RANGE_ONLY_RE.match(line):
            if not current_exp.get("dates"):
                current_exp["dates"] = line.strip()
            continue

        # Check if this is a new job entry (must not be a bullet line)
        if (not is_bullet_line) and _looks_like_job_header(line):
            if current_exp:
                experiences.append(current_exp)
            current_exp = _parse_job_header(line)
        elif current_exp:
            # This is a bullet point
            bullet = _strip_bullet_prefix(line)
            if bullet:
                # Merge "continuation lines" produced by PDF/DOCX text extraction.
                if (
                    (not is_bullet_line)
                    and current_exp.get("bullets")
                    and len(bullet) <= 60
                    and re.match(r"^(and|with|to|for|per|including|including:)\b", bullet, flags=re.I)
                ):
                    current_exp["bullets"][-1] = (current_exp["bullets"][-1].rstrip() + " " + bullet).strip()
                elif (not is_bullet_line) and current_exp.get("bullets") and bullet[:1].islower():
                    current_exp["bullets"][-1] = (current_exp["bullets"][-1].rstrip() + " " + bullet).strip()
                else:
                    current_exp['bullets'].append(bullet)
    
    if current_exp:
        experiences.append(current_exp)
    
    return experiences


def _looks_like_job_header(line: str) -> bool:
    """Check if line looks like a job title/company header."""
    if not line:
        return False
    t = line.strip()
    if not t:
        return False
    # Reject bullet lines (common source of mis-parses)
    if _BULLET_PREFIX_RE.match(t) or t.startswith("•"):
        return False
    # Reject obvious URL lines
    if t.lower().startswith(("url:", "http://", "https://")):
        return False

    # Contains date patterns like "2023 – Present", "Jan 2022 – Dec 2023"
    date_patterns = [
        r'\b20\d{2}\b',  # Year
        r'\bpresent\b',
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b',
    ]
    line_lower = line.lower()
    has_date = any(re.search(p, line_lower) for p in date_patterns)

    # Company separators are typical in headers.
    has_separator = (
        ("|" in t)
        or ("—" in t)
        or ("–" in t)
        or (" - " in t)
        or (" @ " in t)
        or (" at " in line_lower)
    )

    # Avoid treating sentences as headers
    if t.endswith((".", ";")):
        return False

    words = [w for w in re.split(r"\s+", re.sub(r"[^\w\s&/+\-]", " ", t)) if w]
    has_enough_words = len(words) >= 2
    return has_date or (has_separator and has_enough_words and len(t) < 180)


def _parse_job_header(line: str) -> Dict[str, Any]:
    """Parse a job header line into structured data."""
    result = {
        "title": "",
        "company": "",
        "location": "",
        "dates": "",
        "bullets": []
    }
    
    # Remove markers
    line = line.strip('⚙►■●•*- ')
    
    # Try to extract dates
    date_match = re.search(r'(\w+\s+20\d{2}|\b20\d{2})\s*[–—-]\s*(\w+\s+20\d{2}|Present|\b20\d{2})', line, re.IGNORECASE)
    if date_match:
        result['dates'] = date_match.group().strip()
        line = line.replace(date_match.group(), '').strip()
    
    # Split by common separators
    separators = ['|', '—', '–', ' - ']
    parts = [line]
    for sep in separators:
        new_parts = []
        for part in parts:
            new_parts.extend(part.split(sep))
        parts = [p.strip() for p in new_parts if p.strip()]
    
    # Assign parts
    if len(parts) >= 2:
        result['title'] = parts[0].strip()
        result['company'] = parts[1].strip()
        if len(parts) >= 3:
            result['location'] = parts[2].strip()
    elif len(parts) == 1:
        result['title'] = parts[0].strip()
    
    return result


def _parse_projects(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse projects section."""
    projects = []
    current_project = None
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith('—') or line.startswith('━'):
            continue
        
        # If this is a URL-only bullet like "• URL: https://..." don't start a new project.
        line_no_bullet = _strip_bullet_prefix(line)
        if current_project and line_no_bullet.lower().startswith(("url:", "http://", "https://")):
            url_match = re.search(r"https?://[^\s]+", line_no_bullet)
            if url_match:
                current_project["url"] = url_match.group()
            continue

        # Check if this is a new project (starts with marker) and not a URL marker
        if line.startswith('⧉') or line.startswith('►') or (line.startswith('•') and not line_no_bullet.lower().startswith("url:")):
            if current_project:
                projects.append(current_project)
            project_name = _strip_leading_markers(line_no_bullet)
            # Check for URL
            url_match = re.search(r'https?://[^\s]+', project_name)
            url = url_match.group() if url_match else None
            if url:
                project_name = project_name.replace(url, '').strip()
            current_project = {
                "name": project_name,
                "description": "",
                "url": url,
                "bullets": []
            }
        elif current_project:
            # This is description or bullet
            if line_no_bullet.lower().startswith('url:') or line_no_bullet.lower().startswith('http'):
                url_match = re.search(r'https?://[^\s]+', line_no_bullet)
                if url_match:
                    current_project['url'] = url_match.group()
            else:
                current_project['bullets'].append(_strip_bullet_prefix(line))
    
    if current_project:
        projects.append(current_project)
    
    return projects


def _parse_education(lines: List[str]) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """Parse education section. Returns (education, certification_like_entries)."""
    education: List[Dict[str, Any]] = []
    cert_like: List[Dict[str, Any]] = []
    for line in lines:
        line = line.strip('•-* ')
        if not line or line.startswith('—'):
            continue

        # Some resumes mix certifications into Education (e.g. "... (Issued Apr 2025)").
        if re.search(r"\bissued\b|\bcertif", line, re.IGNORECASE):
            cert_like.append({"name": line.strip(), "issuer": "", "date": ""})
            continue
        
        entry = {"degree": "", "school": "", "dates": ""}
        
        # Extract dates
        date_match = re.search(r'\(?(\d{4})\s*[–—-]\s*(\d{4}|Present)?\)?', line, re.IGNORECASE)
        if date_match:
            entry['dates'] = date_match.group().strip('()')
            line = line.replace(date_match.group(), '').strip()
        
        # Split by common separators
        if '—' in line:
            parts = line.split('—')
        elif ' - ' in line:
            parts = line.split(' - ')
        else:
            parts = [line]
        
        if len(parts) >= 2:
            entry['degree'] = parts[0].strip()
            entry['school'] = parts[1].strip()
        else:
            entry['degree'] = line.strip()
        
        if entry['degree']:
            education.append(entry)
    
    return education, cert_like


def _parse_certifications(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse certifications section."""
    certs = []
    for line in lines:
        line = line.strip('•-* ')
        if not line or line.startswith('—'):
            continue
        
        entry = {"name": "", "issuer": "", "date": ""}
        
        # Extract date
        date_match = re.search(r'\(?(Issued\s+)?\w+\s+20\d{2}\)?', line, re.IGNORECASE)
        if date_match:
            entry['date'] = date_match.group().strip('()')
            line = line.replace(date_match.group(), '').strip()
        
        # Split by separator
        if '—' in line:
            parts = line.split('—')
            entry['name'] = parts[0].strip()
            if len(parts) > 1:
                entry['issuer'] = parts[1].strip()
        else:
            entry['name'] = line.strip()
        
        if entry['name']:
            certs.append(entry)
    
    return certs


def _parse_bullets(lines: List[str]) -> List[str]:
    """Parse a list of bullet points."""
    bullets = []
    for line in lines:
        line = line.strip('•-* ')
        if line and not line.startswith('—'):
            bullets.append(line)
    return bullets


# ==================== DOCX Generation ====================

def _clean_docx_text(s: str) -> str:
    t = (s or "").strip()
    if not t:
        return ""
    # ATS-safe punctuation normalization (avoid "smart" punctuation that some parsers mangle)
    t = (
        t.replace("\u2011", "-")  # non-breaking hyphen
        .replace("\u2010", "-")  # hyphen
        .replace("\u2013", "-")  # en dash
        .replace("\u2014", "-")  # em dash
        .replace("\u2212", "-")  # minus
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201C", '"')
        .replace("\u201D", '"')
        .replace("\u00A0", " ")  # NBSP
    )
    # Replace decorative bullets with a safe separator
    t = t.replace("•", " | ")
    # Remove common markdown / prefix noise from LLM outputs
    t = re.sub(r"^\s*(#+\s*)", "", t)
    t = t.replace("**", "").replace("__", "")
    # Collapse weird spacing
    t = re.sub(r"[ \t]{2,}", " ", t)
    return t.strip()


def _split_paragraphs(text: str) -> List[str]:
    t = (text or "").strip()
    if not t:
        return []
    # Normalize newlines
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    parts = [p.strip() for p in re.split(r"\n{2,}", t) if p.strip()]
    # If user/LLM returned single-line with newline bullets, treat each line as its own paragraph
    if len(parts) == 1 and "\n" in parts[0]:
        parts = [ln.strip() for ln in parts[0].split("\n") if ln.strip()]
    return [_clean_docx_text(p) for p in parts if _clean_docx_text(p)]


def _setup_resume_styles(doc: "Document") -> None:
    """
    Normalize the small set of styles we rely on.

    Important: keep it ATS-friendly (no tables required for structure).
    """
    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # We intentionally do NOT rely on Word's built-in list styles for bullets.
    # Different Word templates render those differently, causing indentation drift.

    # ---------- Premium ATS-safe custom styles ----------
    # (Still single-column, no tables/columns/icons; just typographic hierarchy)
    try:
        if "ResumeName" not in [s.name for s in doc.styles]:
            s_name = doc.styles.add_style("ResumeName", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_name = doc.styles["ResumeName"]
        s_name.base_style = doc.styles["Normal"]
        s_name.font.name = "Calibri"
        s_name.font.size = Pt(20)
        s_name.font.bold = True
        s_name.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        s_name.paragraph_format.space_after = Pt(2)
    except Exception:
        pass

    try:
        if "ResumeTitle" not in [s.name for s in doc.styles]:
            s_title = doc.styles.add_style("ResumeTitle", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_title = doc.styles["ResumeTitle"]
        s_title.base_style = doc.styles["Normal"]
        s_title.font.name = "Calibri"
        s_title.font.size = Pt(11.5)
        s_title.font.bold = False
        s_title.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        s_title.paragraph_format.space_after = Pt(4)
    except Exception:
        pass

    try:
        if "ResumeContact" not in [s.name for s in doc.styles]:
            s_contact = doc.styles.add_style("ResumeContact", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_contact = doc.styles["ResumeContact"]
        s_contact.base_style = doc.styles["Normal"]
        s_contact.font.name = "Calibri"
        s_contact.font.size = Pt(9.5)
        s_contact.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        s_contact.paragraph_format.space_after = Pt(10)
    except Exception:
        pass

    try:
        if "ResumeSection" not in [s.name for s in doc.styles]:
            s_sec = doc.styles.add_style("ResumeSection", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_sec = doc.styles["ResumeSection"]
        s_sec.base_style = doc.styles["Normal"]
        s_sec.font.name = "Calibri"
        s_sec.font.size = Pt(11)
        s_sec.font.bold = True
        s_sec.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        s_sec.paragraph_format.space_before = Pt(12)
        s_sec.paragraph_format.space_after = Pt(4)
    except Exception:
        pass

    try:
        if "ResumeBody" not in [s.name for s in doc.styles]:
            s_body = doc.styles.add_style("ResumeBody", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_body = doc.styles["ResumeBody"]
        s_body.base_style = doc.styles["Normal"]
        s_body.font.name = "Calibri"
        s_body.font.size = Pt(10.5)
        s_body.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        s_body.paragraph_format.space_after = Pt(3)
        s_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    except Exception:
        pass

    try:
        if "ResumeBullet" not in [s.name for s in doc.styles]:
            s_b = doc.styles.add_style("ResumeBullet", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_b = doc.styles["ResumeBullet"]
        s_b.base_style = doc.styles["Normal"]
        s_b.font.name = "Calibri"
        s_b.font.size = Pt(10.5)
        s_b.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        pf = s_b.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Inches(0.30)
        pf.first_line_indent = Inches(-0.15)
    except Exception:
        pass

    try:
        if "ResumeRole" not in [s.name for s in doc.styles]:
            s_role = doc.styles.add_style("ResumeRole", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_role = doc.styles["ResumeRole"]
        s_role.base_style = doc.styles["Normal"]
        s_role.font.name = "Calibri"
        s_role.font.size = Pt(11)
        s_role.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        s_role.paragraph_format.space_before = Pt(6)
        s_role.paragraph_format.space_after = Pt(2)
    except Exception:
        pass

    try:
        if "ResumeSubtle" not in [s.name for s in doc.styles]:
            s_sub = doc.styles.add_style("ResumeSubtle", WD_STYLE_TYPE.PARAGRAPH)
        else:
            s_sub = doc.styles["ResumeSubtle"]
        s_sub.base_style = doc.styles["Normal"]
        s_sub.font.name = "Calibri"
        s_sub.font.size = Pt(10)
        s_sub.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        s_sub.paragraph_format.space_after = Pt(3)
    except Exception:
        pass


def _add_bullet_paragraph(doc: "Document", text: str) -> None:
    """
    Add a bullet with explicit hanging indentation (stable across Word templates).
    Uses a hyphen bullet for maximum ATS compatibility.
    """
    t = _strip_bullet_prefix(_strip_leading_markers(_clean_docx_text(text)))
    if not t:
        return
    # Use our deterministic bullet style if available
    try:
        para = doc.add_paragraph(style="ResumeBullet")
    except Exception:
        para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Inches(0.30)
    pf.first_line_indent = Inches(-0.15)  # hanging indent for "- "
    para.add_run(f"- {t}")


def _fix_bullet_verb_tense(text: str) -> str:
    """
    Tiny polish for common present→past tense mistakes that hurt reviewer impression.
    Keep conservative; never invent content.
    """
    t = _clean_docx_text(text)
    if t.lower().startswith("lead "):
        return "Led " + t[5:]
    return t


def _polish_experience_bullet(text: str) -> str:
    """
    Best-effort polish for reviewer/ATS readability.
    Must remain truthful: only reformat/normalize the existing content.
    """
    t = _fix_bullet_verb_tense(text)
    t = _clean_docx_text(t)
    # Expand common abbreviations (keep conservative)
    t = re.sub(r"\bhrs\b", "hours", t, flags=re.IGNORECASE)
    t = re.sub(r"\bhr\b", "hour", t, flags=re.IGNORECASE)
    # Make common metric phrases read naturally (still the same fact)
    t = re.sub(r"(?<!\w)\+(\d+%)\s+conversion\b", r"improving conversion by \1", t, flags=re.IGNORECASE)
    # Remove leading "+" before numeric metrics (e.g. "+34%") while preserving the metric
    t = re.sub(r"(?<!\w)\+(\d)", r"\1", t)
    # Normalize spaces around slashes
    t = re.sub(r"\s*/\s*", "/", t)
    # Ensure consistent ending punctuation
    if t and t[-1] not in ".!?":
        t += "."
    return t


def _add_right_aligned_dates_run(para, dates: str, *, section) -> None:
    """
    Add a right-aligned dates run on the same line using a real tab stop.
    This is ATS-friendly (vs. tables) and avoids alignment hacks.
    """
    d = _clean_docx_text(dates)
    if not d:
        return
    try:
        width = section.page_width - section.left_margin - section.right_margin
        para.paragraph_format.tab_stops.add_tab_stop(width, alignment=WD_TAB_ALIGNMENT.RIGHT)
    except Exception:
        # Best-effort: if tab stops fail, still add dates inline.
        pass
    para.add_run("\t")
    r = para.add_run(d)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.italic = True


def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(margin_value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def _add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    # ATS-minimal mode: avoid decorative lines (some parsers misread them)
    return

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Create bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4A4A4A')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_header(doc, title: str):
    """Add a styled section header."""
    try:
        para = doc.add_paragraph(style="ResumeSection")
    except Exception:
        para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)

    run = para.add_run(_clean_docx_text(title).upper())
    run.bold = True
    run.font.size = Pt(11)
    try:
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    except Exception:
        pass

    # Subtle rule under section title (ATS-safe, still plain text)
    try:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'B0B0B0')
        pBdr.append(bottom)
        pPr.append(pBdr)
    except Exception:
        pass


def _flatten_core_skills(skills: Any) -> List[str]:
    items: List[str] = []
    if isinstance(skills, list):
        for g in skills:
            if isinstance(g, dict) and isinstance(g.get("items"), list):
                items.extend([str(x).strip() for x in g.get("items", []) if str(x).strip()])
            elif isinstance(g, str):
                items.append(g.strip())
    # de-dupe, preserve order
    out: List[str] = []
    seen = set()
    for s in items:
        if not s:
            continue
        key = s.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(s)
    return out


def _add_core_skills(doc: "Document", skills: List[str]) -> None:
    if not skills:
        return
    # 2–3 lines max; comma-separated for clean reading
    per_line = 8
    lines = []
    for i in range(0, min(len(skills), 18), per_line):
        lines.append(", ".join(skills[i : i + per_line]))
    for line in lines[:3]:
        try:
            para = doc.add_paragraph(line, style="ResumeBody")
        except Exception:
            para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(2)


def _select_relevant_keywords(
    *,
    job_keywords: Optional[list[str]],
    resume_text: str,
    tailored_summary: str,
    tailored_bullets: list[Dict[str, Any]],
    max_items: int = 12,
) -> list[str]:
    """
    Select job keywords that already appear in the resume/pack text.
    This improves ATS matching without keyword stuffing or inventing skills.
    """
    if not job_keywords:
        return []
    hay = " ".join(
        [
            (resume_text or ""),
            (tailored_summary or ""),
            " ".join(
                [
                    (b.get("text", "") if isinstance(b, dict) else str(b))
                    for b in (tailored_bullets or [])
                ]
            ),
        ]
    ).lower()
    out: list[str] = []
    seen = set()
    for kw in job_keywords:
        k = _clean_docx_text(str(kw or ""))
        if not k:
            continue
        # Normalize common ATS acronyms
        if k.lower() == "api":
            k = "API"
        elif k.lower() == "rest":
            k = "REST"
        elif k.lower() == "sql":
            k = "SQL"
        elif k.lower() == "ci/cd":
            k = "CI/CD"
        kl = k.lower()
        if kl in seen:
            continue
        def _present(h: str, needle: str) -> bool:
            # Avoid false positives like "Java" matching inside "Javaid".
            n = (needle or "").strip()
            if not n:
                return False
            nl = n.lower()
            # If keyword is a simple word, require word boundaries.
            if re.fullmatch(r"[a-z0-9]+", nl):
                return bool(re.search(rf"(?<![a-z0-9]){re.escape(nl)}(?![a-z0-9])", h, flags=re.I))
            # If keyword contains punctuation/spaces, fall back to substring match.
            return nl in h

        # Only include if it's already present somewhere (no stuffing).
        if _present(hay, kl):
            out.append(k)
            seen.add(kl)
        if len(out) >= max_items:
            break
    return out


def generate_resume_docx(
    tailored_summary: str,
    tailored_bullets: list[Dict[str, Any]],
    original_resume_text: Optional[str] = None,
    job_keywords: Optional[list[str]] = None,
    experience_override: Optional[list[Dict[str, Any]]] = None,
    resume_structure_override: Optional[Dict[str, Any]] = None,
) -> BytesIO:
    """
    Generate a complete ATS-friendly tailored resume DOCX.
    
    Creates a professionally formatted resume with:
    - Proper header with name and contact info
    - Tailored summary section
    - Tailored key achievements
    - Original experience, education, skills formatted properly
    
    Returns:
        BytesIO object containing the DOCX file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    _setup_resume_styles(doc)
    
    # Parse original resume if provided
    parsed = _parse_resume_into_structure(original_resume_text) if original_resume_text else {}
    if isinstance(resume_structure_override, dict) and resume_structure_override:
        # Allow a fully-structured resume override (AI-crafted). Keep it best-effort.
        parsed = {**parsed, **resume_structure_override}
    
    # ==================== HEADER ====================
    # Name
    name = _clean_docx_text(parsed.get('name', 'Your Name'))
    try:
        name_para = doc.add_paragraph(style="ResumeName")
    except Exception:
        name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_para.add_run(name.upper())
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    name_para.paragraph_format.space_after = Pt(2)
    
    # Title/Headline
    title = _clean_docx_text(parsed.get('title', ''))
    if title:
        try:
            title_para = doc.add_paragraph(style="ResumeTitle")
        except Exception:
            title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        title_para.paragraph_format.space_after = Pt(4)
    
    # Contact info (single line, left-aligned)
    contact_items = parsed.get('contact', [])
    location = parsed.get('location', '')
    
    contact_parts = []
    # ATS-friendly ordering: location, phone, email, links
    email = None
    phone = None
    links: list[str] = []
    for item in contact_items:
        t = item.get("type")
        if t == "email" and not email:
            email = item.get("value")
        elif t == "phone" and not phone:
            phone = item.get("value")
        elif t in ("linkedin", "github", "website"):
            links.append(item.get("url", item.get("value", "")))
    if location:
        contact_parts.append(_clean_docx_text(location))
    if phone:
        contact_parts.append(_clean_docx_text(phone))
    if email:
        contact_parts.append(_clean_docx_text(email))
    for u in links:
        u = _clean_docx_text(u)
        if u:
            contact_parts.append(u)
    
    if contact_parts:
        try:
            contact_para = doc.add_paragraph(style="ResumeContact")
        except Exception:
            contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Keep only the most important items to avoid wrapping/indent glitches in Word.
        # Order is already: location, phone, email, links...
        contact_text = " | ".join(contact_parts[:4])
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        contact_para.paragraph_format.space_after = Pt(8)
    # No separate horizontal line; section headers add subtle rule
    
    # ==================== PROFESSIONAL SUMMARY ====================
    _add_section_header(doc, 'PROFESSIONAL SUMMARY')
    
    # Prefer AI-crafted structured summary if provided via override.
    if isinstance(resume_structure_override, dict) and resume_structure_override.get("summary"):
        summary_text = str(resume_structure_override.get("summary") or "")
    else:
        summary_text = tailored_summary if tailored_summary else parsed.get('summary', '')
    if summary_text:
        for ptxt in _split_paragraphs(summary_text)[:4]:
            try:
                summary_para = doc.add_paragraph(_clean_docx_text(ptxt), style="ResumeBody")
            except Exception:
                summary_para = doc.add_paragraph(_clean_docx_text(ptxt))
            summary_para.paragraph_format.space_after = Pt(3)
            summary_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Optional 1-line proof line (years, scope, impact) if provided by optimized resume
        if isinstance(resume_structure_override, dict) and resume_structure_override.get("summary_proof"):
            try:
                proof_para = doc.add_paragraph(_clean_docx_text(resume_structure_override.get("summary_proof")), style="ResumeSubtle")
            except Exception:
                proof_para = doc.add_paragraph(_clean_docx_text(resume_structure_override.get("summary_proof")))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    
    # ==================== CORE SKILLS ====================
    _add_section_header(doc, 'CORE SKILLS')
    core_skills = parsed.get('core_skills')
    skills = core_skills if isinstance(core_skills, list) else parsed.get('skills', [])
    core = _flatten_core_skills(skills)
    _add_core_skills(doc, core)

    # Optional ATS keyword line (only includes keywords already present in resume/pack text)
    kws = _select_relevant_keywords(
        job_keywords=job_keywords,
        resume_text=original_resume_text or "",
        tailored_summary=tailored_summary or "",
        tailored_bullets=tailored_bullets or [],
        max_items=12,
    )
    already_has_kw_line = "relevant keywords:" in (original_resume_text or "").lower()
    if kws and not already_has_kw_line:
        try:
            kw_para = doc.add_paragraph(style="ResumeSubtle")
        except Exception:
            kw_para = doc.add_paragraph()
        r1 = kw_para.add_run("Relevant keywords: ")
        r1.bold = True
        kw_para.add_run(", ".join(kws))
        kw_para.paragraph_format.space_after = Pt(4)
    
    # ==================== EXPERIENCE ====================
    # Prefer full structured experience if provided by optimized_resume.
    override_exp = None
    if isinstance(resume_structure_override, dict) and isinstance(resume_structure_override.get("experience"), list):
        override_exp = resume_structure_override.get("experience")
    experience = override_exp if override_exp else (experience_override if (isinstance(experience_override, list) and experience_override) else parsed.get('experience', []))
    if experience:
        _add_section_header(doc, 'PROFESSIONAL EXPERIENCE')
        
        for exp in experience:
            try:
                job_para = doc.add_paragraph(style="ResumeRole")
            except Exception:
                job_para = doc.add_paragraph()

            title = _clean_docx_text(exp.get("title", "Position"))
            company = _clean_docx_text(exp.get("company", ""))
            location = _clean_docx_text(exp.get("location", ""))
            dates = _clean_docx_text(exp.get("dates", ""))

            left = title
            if company:
                left = f"{left} - {company}" if left else company
            if location:
                left = f"{left}, {location}" if left else location
            header = left
            if dates:
                header = f"{left} | {dates}" if left else dates

            run = job_para.add_run(header)
            run.bold = True
            run.font.size = Pt(11)

            scope_line = exp.get("scope") or exp.get("scope_line")
            if scope_line:
                try:
                    scope_para = doc.add_paragraph(_clean_docx_text(str(scope_line)), style="ResumeSubtle")
                except Exception:
                    scope_para = doc.add_paragraph(_clean_docx_text(str(scope_line)))
                scope_para.paragraph_format.space_after = Pt(2)
            
            # Bullets
            for bullet in exp.get('bullets', []):
                if bullet:
                    _add_bullet_paragraph(doc, _polish_experience_bullet(str(bullet)))
    
    # ==================== PROJECTS ====================
    projects = parsed.get('projects', [])
    if projects:
        _add_section_header(doc, 'PROJECTS')
        
        for proj in projects:
            try:
                proj_header = doc.add_paragraph(style="ResumeRole")
            except Exception:
                proj_header = doc.add_paragraph()
            proj_header.paragraph_format.space_before = Pt(6)
            proj_header.paragraph_format.space_after = Pt(2)
            
            pname = proj.get('name', 'Project')
            if not pname:
                pname = "Project"
            name_run = proj_header.add_run(_clean_docx_text(_strip_leading_markers(pname)))
            name_run.bold = True
            name_run.font.size = Pt(11)

            stack = _clean_docx_text(str(proj.get("stack") or proj.get("tech_stack") or ""))
            if stack.strip().lower() in {"n/a", "na", "none", "null", "unknown"}:
                stack = ""
            url = _clean_docx_text(str(proj.get('url', '') or ""))
            if stack:
                proj_header.add_run(f" - {stack}")
            if url:
                proj_header.add_run(f" | {url}")

            # One description line (not bold) to separate the project clearly
            desc = proj.get("description")
            if desc:
                try:
                    desc_para = doc.add_paragraph(_clean_docx_text(str(desc)), style="ResumeBody")
                except Exception:
                    desc_para = doc.add_paragraph(_clean_docx_text(str(desc)))
                desc_para.paragraph_format.space_after = Pt(2)
            
            for bullet in proj.get('bullets', []):
                if bullet:
                    _add_bullet_paragraph(doc, _clean_docx_text(str(bullet)))

            # Add a small spacer between projects for clarity
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)
    
    # ==================== EDUCATION ====================
    education = parsed.get('education', [])
    if education:
        _add_section_header(doc, 'EDUCATION')
        
        for edu in education:
            try:
                edu_para = doc.add_paragraph(style="ResumeBody")
            except Exception:
                edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(4)
            
            degree = _clean_docx_text(edu.get('degree', ''))
            school = _clean_docx_text(edu.get('school', ''))
            dates = _clean_docx_text(edu.get('dates', ''))
            left = degree
            if school:
                left = f"{left} - {school}" if left else school
            if dates:
                left = f"{left} | {dates}" if left else dates
            run = edu_para.add_run(left)
            run.bold = True
    
    # ==================== CERTIFICATIONS ====================
    certifications = parsed.get('certifications', [])
    if certifications:
        _add_section_header(doc, 'CERTIFICATIONS')
        
        for cert in certifications:
            parts = [_clean_docx_text(cert.get("name", ""))]
            issuer = _clean_docx_text(cert.get("issuer", ""))
            date = _clean_docx_text(cert.get("date", ""))
            if issuer:
                parts.append(issuer)
            if date:
                parts.append(date)
            try:
                cert_para = doc.add_paragraph(style="ResumeBody")
            except Exception:
                cert_para = doc.add_paragraph()
            cert_para.add_run(" - ".join([p for p in parts if p]))
    
    # ==================== ADDITIONAL ====================
    additional = parsed.get('additional', [])
    if additional:
        _add_section_header(doc, 'ADDITIONAL')
        for item in additional:
            if item:
                try:
                    doc.add_paragraph(_clean_docx_text(str(item)), style="ResumeBody")
                except Exception:
                    doc.add_paragraph(_clean_docx_text(str(item)))
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_cover_note_docx(
    cover_note: str,
    job_title: Optional[str] = None,
    company_name: Optional[str] = None,
    applicant_name: Optional[str] = None,
    applicant_email: Optional[str] = None,
    applicant_phone: Optional[str] = None,
    applicant_location: Optional[str] = None,
    applicant_linkedin: Optional[str] = None,
) -> BytesIO:
    """
    Generate a professionally formatted cover letter DOCX.
    
    Includes applicant's contact information in header - ready to use with no edits needed.
    
    Returns:
        BytesIO object containing the DOCX file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # ==================== APPLICANT HEADER ====================
    # Name
    name = applicant_name or 'Applicant'
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_para.paragraph_format.space_after = Pt(2)
    
    # Contact info line
    contact_parts = []
    if applicant_location:
        contact_parts.append(applicant_location)
    if applicant_email:
        contact_parts.append(applicant_email)
    if applicant_phone:
        contact_parts.append(applicant_phone)
    if applicant_linkedin:
        contact_parts.append(applicant_linkedin)
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = '  |  '.join(contact_parts)
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        contact_para.paragraph_format.space_after = Pt(18)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ==================== DATE ====================
    from datetime import datetime
    date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
    date_para.paragraph_format.space_after = Pt(18)
    
    # ==================== RECIPIENT (Company) ====================
    if company_name:
        company_para = doc.add_paragraph()
        company_para.add_run(company_name)
        company_para.paragraph_format.space_after = Pt(2)
    
    # Subject line (if provided)
    if job_title:
        subject_para = doc.add_paragraph()
        subject_para.add_run('Re: Application for ')
        subject_run = subject_para.add_run(job_title)
        subject_run.bold = True
        subject_para.paragraph_format.space_after = Pt(18)
    
    # ==================== GREETING ====================
    greeting_para = doc.add_paragraph('Dear Hiring Manager,')
    greeting_para.paragraph_format.space_after = Pt(12)
    
    # ==================== BODY ====================
    # Cover note content - split into paragraphs
    paragraphs = cover_note.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            # Skip if it's a greeting or closing we'll add ourselves
            if para_text.lower().startswith('dear') or para_text.lower().startswith('sincerely'):
                continue
            # Also skip if it looks like a signature line
            if para_text.startswith('[') and para_text.endswith(']'):
                continue
            para = doc.add_paragraph(para_text)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15
    
    # ==================== CLOSING ====================
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_before = Pt(12)
    closing_para.add_run('Sincerely,')
    
    # Signature with actual name
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_before = Pt(24)
    sig_run = sig_para.add_run(name)
    sig_run.bold = True
    
    # Add email below signature
    if applicant_email:
        email_para = doc.add_paragraph()
        email_para.add_run(applicant_email)
        email_para.paragraph_format.space_before = Pt(2)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_resume_plain_text(
    tailored_summary: str,
    tailored_bullets: list[Dict[str, Any]],
    original_resume_text: Optional[str] = None,
    job_keywords: Optional[list[str]] = None,
    experience_override: Optional[list[Dict[str, Any]]] = None,
    resume_structure_override: Optional[Dict[str, Any]] = None,
) -> str:
    """
    Generate a plain-text version of the resume for ATS preview.
    """
    parsed = _parse_resume_into_structure(original_resume_text) if original_resume_text else {}
    if isinstance(resume_structure_override, dict) and resume_structure_override:
        parsed = {**parsed, **resume_structure_override}

    lines: List[str] = []

    name = _clean_docx_text(parsed.get("name", "Your Name")).upper()
    title = _clean_docx_text(parsed.get("title", ""))
    contact_items = parsed.get("contact", [])
    location = _clean_docx_text(parsed.get("location", ""))

    email = None
    phone = None
    links: list[str] = []
    for item in contact_items:
        t = item.get("type")
        if t == "email" and not email:
            email = item.get("value")
        elif t == "phone" and not phone:
            phone = item.get("value")
        elif t in ("linkedin", "github", "website"):
            links.append(item.get("url", item.get("value", "")))

    contact_parts = []
    if location:
        contact_parts.append(location)
    if phone:
        contact_parts.append(_clean_docx_text(phone))
    if email:
        contact_parts.append(_clean_docx_text(email))
    for u in links:
        u = _clean_docx_text(u)
        if u:
            contact_parts.append(u)

    lines.append(name)
    if title:
        lines.append(title)
    if contact_parts:
        lines.append(" | ".join(contact_parts[:4]))

    # Summary
    lines.append("")
    lines.append("PROFESSIONAL SUMMARY")
    summary_text = ""
    if isinstance(resume_structure_override, dict) and resume_structure_override.get("summary"):
        summary_text = str(resume_structure_override.get("summary") or "")
    else:
        summary_text = tailored_summary if tailored_summary else parsed.get("summary", "")
    for ptxt in _split_paragraphs(summary_text)[:4]:
        lines.append(_clean_docx_text(ptxt))
    if isinstance(resume_structure_override, dict) and resume_structure_override.get("summary_proof"):
        lines.append(_clean_docx_text(resume_structure_override.get("summary_proof")))

    # Core skills
    lines.append("")
    lines.append("CORE SKILLS")
    core_skills = parsed.get("core_skills")
    skills = core_skills if isinstance(core_skills, list) else parsed.get("skills", [])
    core = _flatten_core_skills(skills)
    if core:
        per_line = 8
        for i in range(0, min(len(core), 18), per_line):
            lines.append(", ".join(core[i : i + per_line]))

    # Experience
    lines.append("")
    lines.append("PROFESSIONAL EXPERIENCE")
    override_exp = None
    if isinstance(resume_structure_override, dict) and isinstance(resume_structure_override.get("experience"), list):
        override_exp = resume_structure_override.get("experience")
    experience = override_exp if override_exp else (experience_override if (isinstance(experience_override, list) and experience_override) else parsed.get("experience", []))
    for exp in experience:
        title = _clean_docx_text(exp.get("title", "Position"))
        company = _clean_docx_text(exp.get("company", ""))
        loc = _clean_docx_text(exp.get("location", ""))
        dates = _clean_docx_text(exp.get("dates", ""))
        left = title
        if company:
            left = f"{left} - {company}" if left else company
        if loc:
            left = f"{left}, {loc}" if left else loc
        header = left
        if dates:
            header = f"{left} | {dates}" if left else dates
        if header:
            lines.append(header)
        scope_line = exp.get("scope") or exp.get("scope_line")
        if scope_line:
            lines.append(_clean_docx_text(str(scope_line)))
        for bullet in exp.get("bullets", []) or []:
            if bullet:
                lines.append(f"- {_clean_docx_text(str(bullet))}")

    # Projects
    projects = parsed.get("projects", [])
    if projects:
        lines.append("")
        lines.append("PROJECTS")
        for proj in projects:
            pname = _clean_docx_text(_strip_leading_markers(proj.get("name", "Project")))
            stack = _clean_docx_text(str(proj.get("stack") or proj.get("tech_stack") or ""))
            if stack.strip().lower() in {"n/a", "na", "none", "null", "unknown"}:
                stack = ""
            url = _clean_docx_text(str(proj.get("url", "") or ""))
            header = pname
            if stack:
                header = f"{header} - {stack}"
            if url:
                header = f"{header} | {url}"
            lines.append(header)
            desc = proj.get("description")
            if desc:
                lines.append(_clean_docx_text(str(desc)))
            for bullet in proj.get("bullets", []) or []:
                if bullet:
                    lines.append(f"- {_clean_docx_text(str(bullet))}")
            lines.append("")

    # Education
    education = parsed.get("education", [])
    if education:
        lines.append("")
        lines.append("EDUCATION")
        for edu in education:
            degree = _clean_docx_text(edu.get("degree", ""))
            school = _clean_docx_text(edu.get("school", ""))
            dates = _clean_docx_text(edu.get("dates", ""))
            left = degree
            if school:
                left = f"{left} - {school}" if left else school
            if dates:
                left = f"{left} | {dates}" if left else dates
            if left:
                lines.append(left)

    # Certifications
    certs = parsed.get("certifications", [])
    if certs:
        lines.append("")
        lines.append("CERTIFICATIONS")
        for cert in certs:
            parts = [_clean_docx_text(cert.get("name", ""))]
            issuer = _clean_docx_text(cert.get("issuer", ""))
            date = _clean_docx_text(cert.get("date", ""))
            if issuer:
                parts.append(issuer)
            if date:
                parts.append(date)
            line = " - ".join([p for p in parts if p])
            if line:
                lines.append(line)

    additional = parsed.get("additional", [])
    if additional:
        lines.append("")
        lines.append("ADDITIONAL")
        for item in additional:
            if item:
                lines.append(_clean_docx_text(str(item)))

    return "\n".join([l for l in lines if l is not None])


def generate_apply_pack_zip(
    tailored_summary: str,
    tailored_bullets: list[Dict[str, Any]],
    cover_note: str,
    job_title: Optional[str] = None,
    company_name: Optional[str] = None,
    original_resume_text: Optional[str] = None,
) -> BytesIO:
    """
    Generate a ZIP file containing separate resume and cover letter DOCX files.
    
    This is the preferred format for apply packs - users get two ready-to-use documents.
    
    Returns:
        BytesIO object containing the ZIP file with:
        - resume.docx (tailored resume)
        - cover_letter.docx (personalized cover letter)
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    # Parse resume to get applicant info
    parsed = _parse_resume_into_structure(original_resume_text) if original_resume_text else {}
    
    # Extract applicant contact info
    applicant_name = parsed.get('name', '')
    applicant_location = parsed.get('location', '')
    applicant_email = None
    applicant_phone = None
    applicant_linkedin = None
    
    for contact in parsed.get('contact', []):
        if contact.get('type') == 'email' and not applicant_email:
            applicant_email = contact.get('value')
        elif contact.get('type') == 'phone' and not applicant_phone:
            applicant_phone = contact.get('value')
        elif contact.get('type') == 'linkedin' and not applicant_linkedin:
            applicant_linkedin = contact.get('url', contact.get('value'))
    
    # Generate resume DOCX
    resume_buffer = generate_resume_docx(
        tailored_summary=tailored_summary,
        tailored_bullets=tailored_bullets,
        original_resume_text=original_resume_text,
    )
    
    # Generate cover letter DOCX with applicant info
    cover_buffer = generate_cover_note_docx(
        cover_note=cover_note,
        job_title=job_title,
        company_name=company_name,
        applicant_name=applicant_name,
        applicant_email=applicant_email,
        applicant_phone=applicant_phone,
        applicant_location=applicant_location,
        applicant_linkedin=applicant_linkedin,
    )
    
    # Create ZIP file
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Add resume
        resume_buffer.seek(0)
        zf.writestr('resume.docx', resume_buffer.read())
        
        # Add cover letter
        cover_buffer.seek(0)
        zf.writestr('cover_letter.docx', cover_buffer.read())
    
    zip_buffer.seek(0)
    return zip_buffer


# Keep the old function for backwards compatibility, but make it use ZIP
def generate_combined_docx(
    tailored_summary: str,
    tailored_bullets: list[Dict[str, Any]],
    cover_note: str,
    job_title: Optional[str] = None,
    company_name: Optional[str] = None,
    original_resume_text: Optional[str] = None,
) -> BytesIO:
    """
    Generate a ZIP file containing separate resume and cover letter DOCX files.
    
    Note: This function now returns a ZIP file instead of a combined DOCX for better usability.
    
    Returns:
        BytesIO object containing the ZIP file
    """
    return generate_apply_pack_zip(
        tailored_summary=tailored_summary,
        tailored_bullets=tailored_bullets,
        cover_note=cover_note,
        job_title=job_title,
        company_name=company_name,
        original_resume_text=original_resume_text,
    )
