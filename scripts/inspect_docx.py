"""
Small helper script to inspect a .docx file's paragraph formatting.

Usage:
  python scripts/inspect_docx.py "C:\\path\\to\\file.docx"
"""

from __future__ import annotations

import sys
from typing import Optional


def _s(v: Optional[object]) -> str:
    return str(v) if v is not None else "None"


def main() -> int:
    try:
        from docx import Document
    except Exception as e:
        print(f"ERROR: python-docx not available: {e}")
        return 2

    # PowerShell on Windows can default to cp1252; DOCX text often includes Unicode punctuation.
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # py3.7+
    except Exception:
        pass

    if len(sys.argv) < 2:
        print("Usage: python scripts/inspect_docx.py <path-to-docx>")
        return 2

    path = sys.argv[1]
    doc = Document(path)

    print(f"file: {path}")
    print(f"paragraphs: {len(doc.paragraphs)}")
    print("---")
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").replace("\t", "\\t").strip()
        if not text:
            continue
        pf = para.paragraph_format
        style = getattr(para.style, "name", None)
        print(
            f"{i:03d} style={style!s:22} "
            f"li={_s(pf.left_indent):>10} fi={_s(pf.first_line_indent):>10} "
            f"sb={_s(pf.space_before):>10} sa={_s(pf.space_after):>10} "
            f"ls={_s(pf.line_spacing):>10} "
            f"| {text[:200]}"
        )

    # Tables can also break formatting; print a quick summary
    if doc.tables:
        print("---")
        print(f"tables: {len(doc.tables)}")
        for ti, tbl in enumerate(doc.tables[:10]):
            print(f"table[{ti}] rows={len(tbl.rows)} cols={len(tbl.columns)}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

